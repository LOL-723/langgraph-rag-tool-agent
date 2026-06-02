import hashlib
import json
import logging
import math
import re
import sys
import threading
from pathlib import Path

from fastapi import UploadFile

from core.config import settings
from core.runtime import EXPECTED_PYTHON
from schemas.rag import RagChunk, RagSource, RagUploadResponse


SUPPORTED_RAG_SUFFIXES = {".txt", ".md", ".csv", ".json", ".pdf", ".docx"}
MAX_RETRIEVAL_SUB_QUERIES = 3
RECURSIVE_CHUNK_SEPARATORS = ("\n\n", "\n", "。", ".", "！", "!", "？", "?")

RETRIEVAL_PLAN_PROMPT = (
    "You are a retrieval query planner for a vector database. "
    "Review the user's question before embedding. Rewrite the main question for semantic retrieval, "
    "including useful synonyms, implicit meaning, and document-style terms. "
    "Do not split by default. Only create sub_queries when the question contains multiple independent "
    "retrieval intents that need different evidence from the document. "
    "For one clear question, return an empty sub_queries array. "
    "Do not create sub_queries just for keyword expansion, synonyms, wording cleanup, or implicit meaning; "
    "put those changes into rewritten_query. Do not answer the question. "
    "Return one valid JSON object only with this shape: "
    '{"rewritten_query":"...","sub_queries":["..."]}. '
    f"sub_queries must contain at most {MAX_RETRIEVAL_SUB_QUERIES} items."
)


class RagError(ValueError):
    pass


class RagService:
    def __init__(self) -> None:
        self.storage_dir = Path(settings.RAG_STORAGE_DIR)
        self.max_upload_bytes = settings.RAG_MAX_UPLOAD_MB * 1024 * 1024
        self.chunk_size = settings.RAG_CHUNK_SIZE
        self.chunk_overlap = min(settings.RAG_CHUNK_OVERLAP, max(0, self.chunk_size - 1))
        self.min_chunk_size = min(settings.RAG_MIN_CHUNK_SIZE, self.chunk_size)
        self._lock = threading.Lock()
        self._embedding_model = None
        self._rerank_model = None
        self._collection = None

    async def upload(self, file: UploadFile) -> RagUploadResponse:
        filename = Path(file.filename or "").name
        if not filename:
            raise RagError("filename cannot be empty")

        suffix = Path(filename).suffix.lower()
        if suffix not in SUPPORTED_RAG_SUFFIXES:
            allowed = ", ".join(sorted(SUPPORTED_RAG_SUFFIXES))
            raise RagError(f"unsupported file type: {suffix}. allowed: {allowed}")

        content = await file.read()
        if not content:
            raise RagError("uploaded file cannot be empty")
        if len(content) > self.max_upload_bytes:
            raise RagError(f"file too large. max upload size is {settings.RAG_MAX_UPLOAD_MB}MB")

        document_id = hashlib.sha256(content).hexdigest()[:16]
        doc_dir = self.storage_dir / document_id
        doc_dir.mkdir(parents=True, exist_ok=True)
        for stale_file in doc_dir.iterdir():
            if stale_file.is_file() and stale_file.name != "chunks.json":
                stale_file.unlink()
        file_path = doc_dir / self._safe_filename(filename)
        file_path.write_bytes(content)

        text = self.extract_text(file_path)
        if not text.strip():
            raise RagError("no text could be extracted from this file")

        chunk_texts = self.split_into_chunks(text)
        if not chunk_texts:
            raise RagError("no chunks could be created from this file")

        chunk_ids = [f"{document_id}-{index:04d}" for index in range(len(chunk_texts))]
        metadatas = [
            {"document_id": document_id, "filename": filename, "chunk_index": index}
            for index in range(len(chunk_texts))
        ]
        chunks = [
            RagChunk(id=chunk_ids[index], index=index, content=chunk)
            for index, chunk in enumerate(chunk_texts)
        ]
        (doc_dir / "chunks.json").write_text(
            json.dumps([chunk.model_dump() for chunk in chunks], ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        collection = self._collection_or_none()
        if collection is not None:
            embeddings = self.embed_chunks(chunk_texts)
            try:
                collection.delete(where={"document_id": document_id})
            except Exception:
                pass

            try:
                collection.upsert(
                    ids=chunk_ids,
                    embeddings=embeddings,
                    metadatas=metadatas,
                    documents=chunk_texts,
                )
            except Exception as exc:
                if not self._is_embedding_dimension_error(exc):
                    raise
                logging.warning(
                    "Chroma collection embedding dimension mismatch; recreating collection: %s",
                    exc,
                )
                collection = self._reset_collection()
                collection.upsert(
                    ids=chunk_ids,
                    embeddings=embeddings,
                    metadatas=metadatas,
                    documents=chunk_texts,
                )

        return RagUploadResponse(
            document_id=document_id,
            filename=filename,
            chunk_count=len(chunks),
            chunks=chunks,
        )


    # rag_node 的召回入口，返回从向量数据库或本地 chunks 中召回的片段。
    def retrieve_with_mode(
        self,
        query: str,
        top_k: int,
        document_id: str | None = None,
    ) -> tuple[list[RagSource], str, dict[str, object]]:
        # 召回前先生成检索计划：rewritten_query 负责改写原问题，
        # sub_queries 只承载真正独立的多角度检索意图。
        retrieval_plan = self.build_retrieval_plan(query)
        retrieval_queries = self._normalize_retrieval_queries(retrieval_plan, query)

        # Chroma 支持一次查询多个问题向量，所以这里把多个检索 query
        # 从 [query1, query2] 转成 [[向量1], [向量2]] 后统一提交。
        query_embeddings = self.embed_chunks(retrieval_queries)
        collection = self._collection_or_none()
        if collection is None:
            # Chroma 不可用时走本地 chunks.json 召回，保证缺少向量库依赖时 RAG 仍能工作。
            return self._retrieve_from_local_chunks_multi(
                query_embeddings,
                top_k,
                document_id=document_id,
            ), "local", retrieval_plan

        try:
            if document_id:
                results = collection.query(
                    query_embeddings=query_embeddings,
                    n_results=top_k,
                    where={"document_id": document_id},
                )
            else:
                results = collection.query(
                    query_embeddings=query_embeddings,
                    n_results=top_k,
                )
        except Exception as exc:
            if self._is_embedding_dimension_error(exc):
                logging.warning(
                    "Chroma collection embedding dimension mismatch; using local chunk retrieval: %s",
                    exc,
                )
                # collection 维度和新 embedding 维度不一致时，退回本地多向量召回。
                return self._retrieve_from_local_chunks_multi(
                    query_embeddings,
                    top_k,
                    document_id=document_id,
                ), "local", retrieval_plan
            raise
        # results 中 documents、metadatas、ids 都按 query_embeddings 分组；
        # 需要展开所有问题向量的召回结果，再合并同一个 chunk。
        sources = self._sources_from_chroma_results(results)
        return self._dedupe_sources(sources), "chroma", retrieval_plan


    # rag_node 的重排入口，返回相关性最高的 top_k 个片段。
    def rerank(self, query: str, sources: list[RagSource], top_k: int) -> list[RagSource]:
        if not sources:
            return []
        if self.rerank_model is None:
            return sources[:top_k]
        
        # 每个召回片段和问题组成一个待打分 pair。
        pairs = [(query, source.content) for source in sources]
        # 使用 CrossEncoder 对 pair 的相关性打分。
        scores = self.rerank_model.predict(pairs)
        # 按分数倒序排序，保留相关性最高的 top_k 个片段。
        # zip(sources, scores) 将片段和分数组合成元组，例如 (source, 0.9)。
        # key=lambda item: float(item[1]) 表示按元组中的分数排序。
        ranked = sorted(zip(sources, scores), key=lambda item: float(item[1]), reverse=True)
        return [source for source, _ in ranked[:top_k]]


    # 在用户问题向量化前生成检索计划，返回 {"rewritten_query": str, "sub_queries": list[str]}。
    # rewritten_query 是改写后的主检索问题；sub_queries 是独立多角度检索问题，失败时回退原问题。
    def build_retrieval_plan(self, query: str) -> dict[str, object]:
        fallback = {"rewritten_query": query, "sub_queries": []}
        if not query or not query.strip():
            return fallback

        try:
            from openai import OpenAI

            client = OpenAI(
                api_key=settings.DEEPSEEK_API_KEY,
                base_url=settings.DEEPSEEK_BASE_URL,
                timeout=settings.LLM_TIMEOUT,
            )
            response = client.chat.completions.create(
                model=settings.LLM_MODEL,
                messages=[
                    {"role": "system", "content": RETRIEVAL_PLAN_PROMPT},
                    {"role": "user", "content": query},
                ],
                temperature=settings.LLM_TEMPERATURE,
                response_format={"type": "json_object"},
            )
            content = response.choices[0].message.content or "{}"
            data = json.loads(content)
        except Exception as exc:
            logging.warning("Failed to build RAG retrieval plan; using original query: %s", exc)
            return fallback

        if not isinstance(data, dict):
            return fallback

        rewritten_query = data.get("rewritten_query")
        sub_queries = data.get("sub_queries")
        if not isinstance(rewritten_query, str) or not rewritten_query.strip():
            rewritten_query = query
        if not isinstance(sub_queries, list):
            sub_queries = []

        # 只保留非空字符串，并限制子问题数量，避免一次 RAG 召回产生过多候选片段。
        return {
            "rewritten_query": rewritten_query,
            "sub_queries": [
                item
                for item in sub_queries
                if isinstance(item, str) and item.strip()
            ][:MAX_RETRIEVAL_SUB_QUERIES],
        }

    def split_into_chunks(self, text: str) -> list[str]:
        normalized = re.sub(r"\r\n?", "\n", text).strip()
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        chunks = self._recursive_split(normalized, 0)
        return self._add_chunk_overlap([chunk for chunk in chunks if chunk.strip()])

    # 能使用 Chroma 就使用 Chroma；embedding_model 不可用时退回本地 hash embedding。
    # SentenceTransformer 返回 NumPy 数组，这里用 tolist() 转成原生 list。
    def embed_chunks(self, chunks: list[str]) -> list[list[float]]:
        if self.embedding_model is None:
            return [self._hash_embedding(chunk) for chunk in chunks]
        embeddings = self.embedding_model.encode(chunks, normalize_embeddings=True)
        return embeddings.tolist()

    def extract_text(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix in {".txt", ".md", ".csv", ".json"}:
            return self._read_text(file_path)
        if suffix == ".pdf":
            return self._extract_pdf(file_path)
        if suffix == ".docx":
            return self._extract_docx(file_path)
        raise RagError(f"unsupported file type: {suffix}")

    @property
    def embedding_model(self):
        if self._embedding_model is None:
            with self._lock:
                if self._embedding_model is None:
                    try:
                        from sentence_transformers import SentenceTransformer

                        self._embedding_model = SentenceTransformer(
                            settings.RAG_EMBEDDING_MODEL,
                            token=settings.HF_TOKEN,
                        )
                    except Exception as exc:
                        logging.warning(
                            "Failed to load embedding model %s; using local hash embeddings: %s",
                            settings.RAG_EMBEDDING_MODEL,
                            exc,
                        )
                        self._embedding_model = False
        if self._embedding_model is False:
            return None
        return self._embedding_model

    @property
    def rerank_model(self):
        if self._rerank_model is None:
            with self._lock:
                if self._rerank_model is None:
                    try:
                        from sentence_transformers import CrossEncoder

                        self._rerank_model = CrossEncoder(
                            settings.RAG_RERANK_MODEL,
                            token=settings.HF_TOKEN,
                        )
                    except Exception as exc:
                        logging.warning(
                            "Failed to load rerank model %s; skipping rerank: %s",
                            settings.RAG_RERANK_MODEL,
                            exc,
                        )
                        self._rerank_model = False
        if self._rerank_model is False:
            return None
        return self._rerank_model

    @property
    def collection(self):
        if self._collection is None:
            with self._lock:
                if self._collection is None:
                    try:
                        import chromadb
                    except ModuleNotFoundError as exc:
                        raise RagError(
                            "chromadb is not installed in the FastAPI runtime. "
                            f"python executable: {sys.executable}. "
                            f"expected executable: {EXPECTED_PYTHON}. "
                            "Start with scripts/dev.ps1 or install dependencies with: "
                            "python -m pip install -r requirements.txt"
                        ) from exc

                    Path(settings.RAG_CHROMA_DIR).mkdir(parents=True, exist_ok=True)
                    client = chromadb.PersistentClient(path=settings.RAG_CHROMA_DIR)
                    self._collection = client.get_or_create_collection(name=settings.RAG_COLLECTION_NAME)
        return self._collection

    def _reset_collection(self):
        import chromadb

        client = chromadb.PersistentClient(path=settings.RAG_CHROMA_DIR)
        try:
            client.delete_collection(name=settings.RAG_COLLECTION_NAME)
        except Exception:
            pass
        self._collection = client.get_or_create_collection(name=settings.RAG_COLLECTION_NAME)
        return self._collection

    @staticmethod
    def _is_embedding_dimension_error(exc: Exception) -> bool:
        message = str(exc).lower()
        return "embedding" in message and "dimension" in message

    def _collection_or_none(self):
        try:
            return self.collection
        except RagError as exc:
            logging.warning("Chroma unavailable; using local chunk retrieval: %s", exc)
            return None

    @staticmethod
    def _normalize_retrieval_queries(
        retrieval_plan: dict[str, object],
        original_query: str,
    ) -> list[str]:
        # 将检索计划整理成最终传给 embed_chunks 的 list[str]。
        # 这里只负责取 rewritten_query、追加 sub_queries、过滤空值并去重。
        # 如果 rewritten_query 无效，就回退 original_query，保证至少有一个主检索 query。
        raw_queries: list[str] = []
        rewritten_query = retrieval_plan.get("rewritten_query")
        sub_queries = retrieval_plan.get("sub_queries")

        if isinstance(rewritten_query, str) and rewritten_query.strip():
            raw_queries.append(rewritten_query)
        else:
            raw_queries.append(original_query)

        if isinstance(sub_queries, list):
            raw_queries.extend(
                item
                for item in sub_queries
                if isinstance(item, str) and item.strip()
            )

        queries: list[str] = []
        seen: set[str] = set()
        for query in raw_queries:
            normalized = query.strip()
            key = normalized.casefold()
            if not normalized or key in seen:
                continue
            seen.add(key)
            queries.append(normalized)
        return queries or [original_query]

    # Chroma 多 query 查询返回二维结果：外层按问题向量分组，内层是该问题的 top_k 片段。
    def _sources_from_chroma_results(self, results: dict[str, object]) -> list[RagSource]:
        documents_groups = results.get("documents") or []
        metadatas_groups = results.get("metadatas") or []
        ids_groups = results.get("ids") or []

        if not isinstance(documents_groups, list):
            return []

        sources: list[RagSource] = []
        for query_index, documents in enumerate(documents_groups):
            # metadatas 和 ids 与 documents 使用相同的分组下标，三者按位置配对。
            metadatas = self._group_at(metadatas_groups, query_index)
            ids = self._group_at(ids_groups, query_index)
            if not isinstance(documents, list):
                continue

            for result_index, content in enumerate(documents):
                metadata = self._item_at(metadatas, result_index) or {}
                chunk_id = self._item_at(ids, result_index)
                if not isinstance(metadata, dict):
                    metadata = {}
                sources.append(
                    RagSource(
                        chunk_id=str(chunk_id or ""),
                        document_id=str(metadata.get("document_id", "")),
                        filename=str(metadata.get("filename", "")),
                        chunk_index=int(metadata.get("chunk_index", 0)),
                        content=str(content or ""),
                    )
                )
        return sources

    @staticmethod
    def _group_at(groups: object, index: int) -> object:
        if isinstance(groups, list) and index < len(groups):
            return groups[index]
        return []

    @staticmethod
    def _item_at(items: object, index: int) -> object:
        if isinstance(items, list) and index < len(items):
            return items[index]
        return None

    @staticmethod
    def _dedupe_sources(sources: list[RagSource]) -> list[RagSource]:
        deduped: list[RagSource] = []
        seen: set[str] = set()
        for source in sources:
            # 多个检索角度可能召回同一个 chunk，优先用 Chroma id 去重。
            key = source.chunk_id.strip()
            if not key:
                key = f"{source.document_id}:{source.chunk_index}"
            if key in seen:
                continue
            seen.add(key)
            deduped.append(source)
        return deduped

    def _retrieve_from_local_chunks_multi(
        self,
        query_embeddings: list[list[float]],
        top_k: int,
        document_id: str | None = None,
    ) -> list[RagSource]:
        sources: list[RagSource] = []
        for query_embedding in query_embeddings:
            # 本地 fallback 仍按单向量计算余弦相似度，再把多角度结果合并去重。
            sources.extend(
                self._retrieve_from_local_chunks(
                    query_embedding,
                    top_k,
                    document_id=document_id,
                )
            )
        return self._dedupe_sources(sources)

    def _retrieve_from_local_chunks(
        self,
        query_embedding: list[float],
        top_k: int,
        document_id: str | None = None,
    ) -> list[RagSource]:
        # Local fallback retrieval path. This intentionally avoids importing or
        # calling Chroma so RAG can still work when chromadb is missing.
        scored_sources: list[tuple[float, RagSource]] = []

        for chunks_file in self.storage_dir.glob("*/chunks.json"):
            current_document_id = chunks_file.parent.name
            if document_id and current_document_id != document_id:
                continue
            try:
                chunks_data = json.loads(chunks_file.read_text(encoding="utf-8"))
            except (OSError, json.JSONDecodeError):
                continue

            source_file = next(
                (path for path in chunks_file.parent.iterdir() if path.name != "chunks.json"),
                None,
            )
            filename = source_file.name if source_file else current_document_id

            valid_chunks: list[tuple[dict[str, object], str]] = []
            for chunk_data in chunks_data:
                if not isinstance(chunk_data, dict):
                    continue
                content = str(chunk_data.get("content", ""))
                if not content:
                    continue
                valid_chunks.append((chunk_data, content))

            if not valid_chunks:
                continue

            chunk_embeddings = self.embed_chunks([content for _, content in valid_chunks])
            for (chunk_data, content), chunk_embedding in zip(valid_chunks, chunk_embeddings):
                score = self._cosine_similarity(query_embedding, chunk_embedding)
                source = RagSource(
                    chunk_id=str(chunk_data.get("id", "")),
                    document_id=current_document_id,
                    filename=filename,
                    chunk_index=int(chunk_data.get("index", 0)),
                    content=content,
                )
                scored_sources.append((score, source))

        scored_sources.sort(key=lambda item: item[0], reverse=True)
        return [source for _, source in scored_sources[:top_k]]

    def _recursive_split(self, text: str, separator_index: int) -> list[str]:
        text = text.strip()
        if not text:
            return []
        if separator_index >= len(RECURSIVE_CHUNK_SEPARATORS):
            if len(text) <= self.chunk_size:
                return [text]
            return self._fixed_length_split(text)

        separator = RECURSIVE_CHUNK_SEPARATORS[separator_index]
        parts = self._split_with_separator(text, separator)
        if len(parts) == 1:
            if len(text) <= self.chunk_size:
                return [text]
            return self._recursive_split(text, separator_index + 1)

        chunks: list[str] = []
        for part in parts:
            if len(part) > self.chunk_size:
                chunks.extend(self._recursive_split(part, separator_index + 1))
                continue
            chunks.append(part)
        return self._merge_small_chunks(chunks)

    @staticmethod
    def _split_with_separator(text: str, separator: str) -> list[str]:
        pieces = text.split(separator)
        parts: list[str] = []
        for index, piece in enumerate(pieces):
            if not piece.strip():
                continue
            part = piece
            if index < len(pieces) - 1:
                part = f"{part}{separator}"
            if part.strip():
                parts.append(part)
        return parts

    def _merge_small_chunks(self, parts: list[str]) -> list[str]:
        chunks: list[str] = []
        for part in parts:
            part = part.strip()
            if not part:
                continue
            if len(part) > self.chunk_size:
                chunks.extend(self._fixed_length_split(part))
                continue

            if not chunks:
                chunks.append(part)
                continue

            previous = chunks[-1]
            candidate = f"{previous}{part}"
            if len(previous) < self.min_chunk_size and len(candidate) <= self.chunk_size:
                chunks[-1] = candidate
            else:
                chunks.append(part)

        if len(chunks) > 1 and len(chunks[-1]) < self.min_chunk_size:
            candidate = f"{chunks[-2]}{chunks[-1]}"
            if len(candidate) <= self.chunk_size:
                chunks[-2] = candidate
                chunks.pop()
        return chunks

    def _fixed_length_split(self, text: str) -> list[str]:
        return [
            text[start:start + self.chunk_size].strip()
            for start in range(0, len(text), self.chunk_size)
            if text[start:start + self.chunk_size].strip()
        ]

    def _add_chunk_overlap(self, chunks: list[str]) -> list[str]:
        if self.chunk_overlap <= 0:
            return chunks

        overlapped: list[str] = []
        for index, chunk in enumerate(chunks):
            room = self.chunk_size - len(chunk)
            if index > 0 and room > 0:
                previous = chunks[index - 1]
                overlap = previous[-min(self.chunk_overlap, room):]
                chunk = f"{overlap}{chunk}"
            overlapped.append(chunk)
        return overlapped

    @staticmethod
    def _read_text(file_path: Path) -> str:
        for encoding in ("utf-8", "utf-8-sig", "gb18030"):
            try:
                return file_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return file_path.read_text(encoding="utf-8", errors="ignore")

    @staticmethod
    def _extract_pdf(file_path: Path) -> str:
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        pages = []
        for page_index, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(f"[page {page_index}]\n{page_text}")
        return "\n\n".join(pages)

    @staticmethod
    def _extract_docx(file_path: Path) -> str:
        from docx import Document

        doc = Document(str(file_path))
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        table_rows = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_rows.append(" | ".join(cells))
        return "\n\n".join(paragraphs + table_rows)

    @staticmethod
    def _safe_filename(filename: str) -> str:
        stem = Path(filename).stem
        suffix = Path(filename).suffix.lower()
        safe_stem = re.sub(r"[^A-Za-z0-9._-]+", "_", stem).strip("._")
        return f"{safe_stem or 'document'}{suffix}"

    @staticmethod
    def _hash_embedding(text: str, dimensions: int = 384) -> list[float]:
        vector = [0.0] * dimensions
        tokens = re.findall(r"[\w\u4e00-\u9fff]+", text.lower())
        if not tokens:
            tokens = [text]
        for token in tokens:
            digest = hashlib.sha256(token.encode("utf-8")).digest()
            for offset in range(0, len(digest), 2):
                index = int.from_bytes(digest[offset:offset + 2], "little") % dimensions
                sign = 1.0 if digest[offset] % 2 == 0 else -1.0
                vector[index] += sign
        norm = math.sqrt(sum(value * value for value in vector)) or 1.0
        return [value / norm for value in vector]

    @staticmethod
    def _cosine_similarity(left: list[float], right: list[float]) -> float:
        return sum(left_value * right_value for left_value, right_value in zip(left, right))


rag_service = RagService()
